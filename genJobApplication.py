# dependencies
import docx 
import re
import os
from docx2pdf import convert
from datetime import date

# set company name
company_name = input("Enter company name:")
position_name = input("Enter position name:")
today = date.today().strftime("%d/%m/%Y")

# set working directory to location of job applications
path = 'C:\\Users\\...\\Word Applications'
os.chdir(path)

# set document to job application template
document = docx.Document('Cover Letter Template.docx')

# name of new application
file_name = position_name + "_" + company_name + ".docx"

# set document to new application
new_doc = document.save(file_name)

# open document
new_doc = docx.Document(file_name)

# function to replace text in template cover letter
def fReplaceText(document, to_replace, replace):
    
    # iterate over paragraphs in document
    for p in document.paragraphs:
        
        # check if text is in paragraph text
        if to_replace.search(p.text):
            
            # find runs (sequence of characters that share same formatting)
            inline = p.runs
            
            # loop over runs
            for i in range(len(inline)):
                if to_replace.search(inline[i].text):
                    text = to_replace.sub(replace, inline[i].text)
                    inline[i].text = text
                   
 # replace text
to_replace1 = re.compile(r"Date")
replace1 = today
to_replace2 = re.compile(r"CompanyName")
replace2 = company_name
to_replace3 = re.compile(r"PositionName")
replace3 = position_name
fReplaceText(new_doc, to_replace1, replace1)
fReplaceText(new_doc, to_replace2, replace2)
fReplaceText(new_doc, to_replace3, replace3)
new_doc.save(file_name)

### convert to pdf
convert(file_name, 'C:\\Users\\...\\PDF Applications')
